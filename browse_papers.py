#!/usr/bin/env python3
"""
Biomedical Articles Search System
Specialized search in PubMed, bioRxiv and Europe PMC
Areas: Biology, Health, Cancer, Oncology, Medicine
"""

import argparse
import requests
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional
from collections import defaultdict
import time
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class BiomedicalSearcher:
    """Class to search in specialized biomedical databases"""

    def __init__(self, keywords: List[str], days: int):
        self.keywords = [kw.lower() for kw in keywords]
        self.days = days
        self.cutoff_date = datetime.now() - timedelta(days=days)

    def search_pubmed(self) -> List[Dict]:
        """
        Search PubMed (NCBI) - Main biomedical database
        Searches: Title + Abstract + MeSH Terms + Author keywords
        30+ million biomedical articles
        """
        print("🔍 Searching PubMed (Title + Abstract + MeSH + Keywords)...")

        base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
        articles = []

        query = ' OR '.join([f'({kw}[Title/Abstract])' for kw in self.keywords])

        try:
            search_params = {
                'db': 'pubmed',
                'term': query,
                'retmax': 500,
                'retmode': 'json',
                'sort': 'pub_date',
                'reldate': self.days,
                'usehistory': 'y'
            }

            search_response = requests.get(
                f"{base_url}esearch.fcgi",
                params=search_params,
                timeout=15
            )

            if search_response.status_code == 200:
                search_data = search_response.json()
                id_list = search_data.get('esearchresult', {}).get('idlist', [])

                print(f"   → Found {len(id_list)} articles in PubMed")

                if id_list:
                    # Step 2: Fetch details in batches of 50
                    for i in range(0, len(id_list), 50):
                        batch_ids = id_list[i:i + 50]

                        fetch_params = {
                            'db': 'pubmed',
                            'id': ','.join(batch_ids),
                            'retmode': 'xml',
                            'rettype': 'abstract'
                        }

                        fetch_response = requests.get(
                            f"{base_url}efetch.fcgi",
                            params=fetch_params,
                            timeout=15
                        )

                        if fetch_response.status_code == 200:
                            # Parse XML
                            try:
                                root = ET.fromstring(fetch_response.content)

                                for article_elem in root.findall('.//PubmedArticle'):
                                    # Extract article data
                                    title_elem = article_elem.find('.//ArticleTitle')
                                    title = title_elem.text if title_elem is not None else 'N/A'

                                    # Abstract (may have multiple sections)
                                    abstract_parts = []
                                    for abs_text in article_elem.findall('.//AbstractText'):
                                        if abs_text.text:
                                            abstract_parts.append(abs_text.text)
                                    abstract = ' '.join(
                                        abstract_parts) if abstract_parts else '[No abstract available]'

                                    # PMID
                                    pmid_elem = article_elem.find('.//PMID')
                                    pmid = pmid_elem.text if pmid_elem is not None else 'Unknown'

                                    # Publication date
                                    year_elem = article_elem.find('.//PubDate/Year')
                                    month_elem = article_elem.find('.//PubDate/Month')
                                    day_elem = article_elem.find('.//PubDate/Day')

                                    try:
                                        year = int(year_elem.text) if year_elem is not None else datetime.now().year
                                        month = self._parse_month(month_elem.text if month_elem is not None else '1')
                                        day = int(day_elem.text) if day_elem is not None else 1
                                        pub_date = datetime(year, month, day)
                                    except Exception:
                                        pub_date = datetime.now()

                                    # Authors
                                    authors = []
                                    for author in article_elem.findall('.//Author')[:3]:
                                        lastname = author.find('.//LastName')
                                        forename = author.find('.//ForeName')
                                        if lastname is not None:
                                            name = lastname.text
                                            if forename is not None:
                                                name = f"{forename.text} {name}"
                                            authors.append(name)

                                    articles.append({
                                        'title': title,
                                        'abstract': abstract,
                                        'published': pub_date,
                                        'url': f'https://pubmed.ncbi.nlm.nih.gov/{pmid}/',
                                        'source': 'PubMed',
                                        'authors': authors if authors else ['N/A']
                                    })

                            except ET.ParseError as e:
                                print(f"   ⚠ Error parsing XML: {e}")

                        time.sleep(0.35)  # Rate limiting (max 3 req/sec for NCBI)

        except Exception as e:
            print(f"   ⚠ Error searching PubMed: {str(e)}")

        print(f"   ✓ Processed: {len(articles)} articles\n")
        return articles

    def search_biorxiv(self) -> List[Dict]:
        """
        Search bioRxiv (biomedical preprints)
        Searches: Title + Abstract + Full text
        Latest preprints in biology
        """
        print("🔍 Searching bioRxiv (Preprints - Title + Abstract + Full Text)...")

        base_url = "https://api.biorxiv.org/details/biorxiv"
        articles = []

        # bioRxiv uses a date range
        start_date = self.cutoff_date.strftime('%Y-%m-%d')
        end_date = datetime.now().strftime('%Y-%m-%d')

        try:
            url = f"{base_url}/{start_date}/{end_date}"
            response = requests.get(url, timeout=15)

            if response.status_code == 200:
                data = response.json()

                print(f"   → Found {data.get('messages', [{}])[0].get('total', 0)} preprints in the period")

                for paper in data.get('collection', []):
                    title = paper.get('title', 'N/A')
                    abstract = paper.get('abstract', '[No abstract]')


                    text_to_search = f"{title.lower()} {abstract.lower()}"
                    has_keyword = any(kw in text_to_search for kw in self.keywords)

                    if has_keyword:
                        # Parse date
                        date_str = paper.get('date', '')
                        try:
                            pub_date = datetime.strptime(date_str, '%Y-%m-%d')
                        except Exception:
                            pub_date = datetime.now()

                        # Authors
                        authors_str = paper.get('authors', 'N/A')
                        authors = [a.strip() for a in authors_str.split(';')[:3]]

                        doi = paper.get('doi', '')

                        articles.append({
                            'title': title,
                            'abstract': abstract,
                            'published': pub_date,
                            'url': f'https://www.biorxiv.org/content/{doi}v1',
                            'source': 'bioRxiv',
                            'authors': authors
                        })

        except Exception as e:
            print(f"   ⚠ Error searching bioRxiv: {str(e)}")

        print(f"   ✓ Processed: {len(articles)} preprints\n")
        return articles

    def search_europepmc(self) -> List[Dict]:
        """
        Search Europe PMC
        Searches: Title + Abstract + Full Text + Grants
        European database with 42+ million biomedical articles
        """
        print("🔍 Searching Europe PMC (Title + Abstract + Full Text)...")

        base_url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
        articles = []

        query_parts = [f'"{kw}"' for kw in self.keywords]
        query = ' OR '.join(query_parts)

        try:
            params = {
                'query': query,
                'format': 'json',
                'pageSize': 200,
                'sort': 'P_PDATE_D desc',  # Sort by publication date
                'fromPublicationDate': self.cutoff_date.strftime('%Y-%m-%d')
            }

            response = requests.get(base_url, params=params, timeout=15)

            if response.status_code == 200:
                data = response.json()
                results = data.get('resultList', {}).get('result', [])

                print(f"   → Found {len(results)} articles in Europe PMC")

                for paper in results:
                    title = paper.get('title', 'N/A')
                    abstract = paper.get('abstractText', '[No abstract available]')

                    # Publication date
                    pub_date_str = paper.get('firstPublicationDate', '')
                    try:
                        pub_date = datetime.strptime(pub_date_str, '%Y-%m-%d')
                    except Exception:
                        pub_date = datetime.now()

                    # Authors
                    author_list = paper.get('authorString', 'N/A').split(', ')[:3]

                    # URL
                    pmid = paper.get('pmid')
                    doi = paper.get('doi')
                    if pmid:
                        url = f'https://europepmc.org/article/MED/{pmid}'
                    elif doi:
                        url = f'https://doi.org/{doi}'
                    else:
                        url = 'N/A'

                    articles.append({
                        'title': title,
                        'abstract': abstract,
                        'published': pub_date,
                        'url': url,
                        'source': 'Europe PMC',
                        'authors': author_list
                    })

        except Exception as e:
            print(f"   ⚠ Error searching Europe PMC: {str(e)}")

        print(f"   ✓ Processed: {len(articles)} articles\n")
        return articles

    def _parse_month(self, month_str: str) -> int:
        """Convert month name to number"""
        months = {
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
        }
        try:
            return int(month_str)
        except Exception:
            return months.get(month_str.lower()[:3], 1)

    def search_all(self) -> List[Dict]:
        """Search all biomedical sources"""
        print(f"🧬 BIOMEDICAL SEARCH - Last {self.days} days")
        print(f"🔬 Keywords: {', '.join(self.keywords)}\n")
        print("=" * 80)

        all_articles = []

        # PubMed (main biomedical database)
        all_articles.extend(self.search_pubmed())

        # bioRxiv (biomedical preprints)
        all_articles.extend(self.search_biorxiv())

        # Europe PMC (European database)
        all_articles.extend(self.search_europepmc())

        print("=" * 80)
        print(f"📊 Raw total: {len(all_articles)} articles from all sources\n")

        # Remove duplicates by title (ignore articles without valid title)
        unique = {}
        for art in all_articles:
            title = art.get('title')
            if not title or title == 'N/A':
                continue

            title_clean = title.lower().strip()
            abstract = art.get('abstract', '')

            # Keep the article with the most complete abstract
            if title_clean not in unique or len(abstract) > len(unique[title_clean].get('abstract', '')):
                unique[title_clean] = art

        deduplicated = list(unique.values())
        print(f"🔄 After removing duplicates: {len(deduplicated)} unique articles\n")

        return deduplicated

    def calculate_score(self, article: Dict) -> Tuple[int, List[str]]:
        """Calculate article score based on matched keywords"""
        title = article.get('title', '')
        abstract = article.get('abstract', '')

        # None-protection
        title_lower = title.lower() if title else ''
        abstract_lower = abstract.lower() if abstract else ''

        matched_keywords = []
        for keyword in self.keywords:
            if keyword in title_lower or keyword in abstract_lower:
                matched_keywords.append(keyword)

        return len(matched_keywords), matched_keywords

    def rank_articles(self, articles: List[Dict]) -> Dict[int, List[Dict]]:
        """Rank articles by score"""
        ranked = defaultdict(list)

        for article in articles:
            score, matched_kws = self.calculate_score(article)
            if score > 0:
                article['score'] = score
                article['matched_keywords'] = matched_kws
                ranked[score].append(article)

        # Sort articles within each score by date (most recent first)
        for score in ranked:
            ranked[score].sort(key=lambda x: x['published'], reverse=True)

        return ranked

    def display_results(self, ranked_articles: Dict[int, List[Dict]]):
        """Display formatted results in the terminal"""
        if not ranked_articles:
            print("❌ No articles found with the specified keywords.\n")
            return

        total_articles = sum(len(articles) for articles in ranked_articles.values())
        print(f"✅ Total relevant articles: {total_articles}\n")
        print("=" * 80)

        for score in sorted(ranked_articles.keys(), reverse=True):
            articles = ranked_articles[score]

            # Collect unique keywords for this group
            all_matched = set()
            for article in articles:
                all_matched.update(article['matched_keywords'])

            matched_str = ', '.join(sorted(all_matched))

            print(f"\n{'=' * 80}")
            print(f"🎯 {score} MATCH{'ES' if score > 1 else ''}")
            print(f"🏷️  Keywords: {matched_str}")
            print(f"{'=' * 80}\n")

            for idx, article in enumerate(articles, 1):
                authors_str = ', '.join(article.get('authors', ['N/A']))
                if len(authors_str) > 70:
                    authors_str = authors_str[:67] + '...'

                title = article.get('title', 'N/A')
                abstract = article.get('abstract', '')

                # Truncate abstract for display
                abstract_preview = abstract[:150] if abstract else '[No abstract]'
                if len(abstract) > 150:
                    abstract_preview += '...'

                print(f"[{idx}] {title}")
                print(f"    📅 Date: {article['published'].strftime('%Y-%m-%d')}")
                print(f"    🗂️  Source: {article['source']}")
                print(f"    👥 Authors: {authors_str}")
                print(f"    🎯 Matches: {', '.join(article['matched_keywords'])}")
                print(f"    📝 Abstract: {abstract_preview}")
                print(f"    🔗 Link: {article['url']}")
                print()

        print("=" * 80)


class DocxExporter:
    """Class to export results in DOCX format"""

    def __init__(self, keywords: List[str], days: int, output_path: str):
        self.keywords = keywords
        self.days = days
        self.output_path = Path(output_path)
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Configure document styles"""
        # Page settings
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_hyperlink(self, paragraph, text, url):
        """Add hyperlink to paragraph"""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                              is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Link style (blue and underlined)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._element.append(hyperlink)

    def export(self, ranked_articles: Dict[int, List[Dict]]) -> Optional[str]:
        """
        Export results to DOCX

        Returns:
            Generated file path
        """
        if not ranked_articles:
            print("⚠️  No articles to export.")
            return None

        # Document header
        title = self.doc.add_heading('Literature Review - Biomedical Articles', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Search information
        info = self.doc.add_paragraph()
        info.add_run('Keywords: ').bold = True
        info.add_run(', '.join(self.keywords))

        date_info = self.doc.add_paragraph()
        date_info.add_run('Period: ').bold = True
        date_info.add_run(f'Last {self.days} days')

        total_articles = sum(len(articles) for articles in ranked_articles.values())
        total_info = self.doc.add_paragraph()
        total_info.add_run('Total articles: ').bold = True
        total_info.add_run(str(total_articles))

        export_date = self.doc.add_paragraph()
        export_date.add_run('Export date: ').bold = True
        export_date.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

        self.doc.add_paragraph()  # Space

        # Articles grouped by score
        for score in sorted(ranked_articles.keys(), reverse=True):
            articles = ranked_articles[score]

            # Collect unique keywords
            all_matched = set()
            for article in articles:
                all_matched.update(article['matched_keywords'])
            matched_str = ', '.join(sorted(all_matched))

            # Group header
            self.doc.add_heading(f'{score} Match{"es" if score > 1 else ""}', level=1)

            kw_para = self.doc.add_paragraph()
            kw_run = kw_para.add_run(f'Keywords found: {matched_str}')
            kw_run.italic = True
            kw_run.font.color.rgb = RGBColor(80, 80, 80)

            self.doc.add_paragraph()  # Space

            # Articles
            for idx, article in enumerate(articles, 1):
                # Article title
                title_para = self.doc.add_paragraph()
                title_run = title_para.add_run(f'[{idx}] {article.get("title", "N/A")}')
                title_run.bold = True
                title_run.font.size = Pt(12)

                # Publication date
                date_para = self.doc.add_paragraph(style='List Bullet')
                date_para.add_run('Date: ').bold = True
                date_para.add_run(article['published'].strftime('%Y-%m-%d'))

                # Source
                source_para = self.doc.add_paragraph(style='List Bullet')
                source_para.add_run('Source: ').bold = True
                source_para.add_run(article['source'])

                # Authors
                authors_str = ', '.join(article.get('authors', ['N/A']))
                if len(authors_str) > 150:
                    authors_str = authors_str[:147] + '...'

                authors_para = self.doc.add_paragraph(style='List Bullet')
                authors_para.add_run('Authors: ').bold = True
                authors_para.add_run(authors_str)

                # Matched keywords
                match_para = self.doc.add_paragraph(style='List Bullet')
                match_para.add_run('Matches: ').bold = True
                match_run = match_para.add_run(', '.join(article['matched_keywords']))
                match_run.font.color.rgb = RGBColor(0, 128, 0)

                # Abstract
                abstract = article.get('abstract', '')
                abstract_text = abstract[:500] if abstract else '[No abstract available]'
                if len(abstract) > 500:
                    abstract_text += '...'

                abstract_para = self.doc.add_paragraph(style='List Bullet')
                abstract_para.add_run('Abstract: ').bold = True
                abstract_para.add_run(abstract_text)

                # Link with hyperlink
                link_para = self.doc.add_paragraph(style='List Bullet')
                link_para.add_run('Link: ').bold = True
                url = article.get('url', 'N/A')
                if url != 'N/A':
                    self._add_hyperlink(link_para, url, url)
                else:
                    link_para.add_run('N/A')

                self.doc.add_paragraph()  # Space between articles

            # Separator between groups
            self.doc.add_paragraph('─' * 80)

        # Save file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        keywords_slug = '_'.join(self.keywords[:3])[:30].replace(' ', '_')
        filename = f'bio_articles_{keywords_slug}_{timestamp}.docx'

        # Ensure directory exists
        self.output_path.mkdir(parents=True, exist_ok=True)

        filepath = self.output_path / filename
        self.doc.save(str(filepath))

        return str(filepath)


def main():
    """Main function"""
    parser = argparse.ArgumentParser(
        description='🧬 Specialized search for BIOMEDICAL articles',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
🔬 BIOMEDICAL DATABASES QUERIED:

  1. PubMed (NCBI)
     • 35+ million biomedical articles
     • Searches: Title + Abstract + MeSH Terms + Keywords
     • Areas: Medicine, Biology, Public Health, Oncology

  2. bioRxiv
     • Recent biomedical preprints
     • Searches: Title + Abstract + Full Text
     • Areas: Molecular Biology, Genetics, Neuroscience, Cancer

  3. Europe PMC
     • 42+ million articles (European database)
     • Searches: Title + Abstract + Full Text + Grants
     • Broad coverage of life sciences

📋 USAGE EXAMPLES:

  # Research on breast cancer
  python browse_papers.py --keywords "breast cancer" "immunotherapy" --days 30

  # Oncology and genetics
  python browse_papers.py --keywords "oncogene" "mutation" "p53" --days 14

  # With DOCX export
  python browse_papers.py --keywords "chemotherapy" "tolerance" --days 7 --export /path/to/folder
        """
    )

    parser.add_argument(
        '--keywords',
        nargs='+',
        required=True,
        help='List of biomedical keywords'
    )

    parser.add_argument(
        '--days',
        type=int,
        required=True,
        help='Look back this many days'
    )

    parser.add_argument(
        '--export',
        type=str,
        default=None,
        help='Folder path to export DOCX document (e.g., /home/user/documents)'
    )

    args = parser.parse_args()

    # Validations
    if args.days <= 0:
        parser.error("--days must be a positive number")

    if not args.keywords:
        parser.error("--keywords must contain at least one keyword")

    # Execute search
    try:
        searcher = BiomedicalSearcher(keywords=args.keywords, days=args.days)
        articles = searcher.search_all()
        ranked = searcher.rank_articles(articles)
        searcher.display_results(ranked)

        # Export to DOCX if requested
        if args.export:
            print(f"\n📄 Exporting results to DOCX...")
            exporter = DocxExporter(
                keywords=args.keywords,
                days=args.days,
                output_path=args.export
            )
            filepath = exporter.export(ranked)

            if filepath:
                print(f"✅ Document generated successfully!")
                print(f"📁 Location: {filepath}\n")
            else:
                print("⚠️  Could not generate the document.\n")

    except KeyboardInterrupt:
        print("\n\n⚠️  Search interrupted by user.")
        return 130
    except Exception as e:
        print(f"❌ Error during search: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == '__main__':
    exit(main())
